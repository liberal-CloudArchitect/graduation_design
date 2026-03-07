import json, os
from docx import Document
from docx.shared import Pt, Emu

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONTENT = os.path.join(SCRIPT_DIR, "task_book_content.json")
FONT_SIZE = Pt(14)
INDENT = Emu(355600)

def fill_cell(cell, paragraphs, indent=True):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    first.clear()
    run = first.add_run(paragraphs[0])
    run.font.size = FONT_SIZE
    if indent:
        first.paragraph_format.first_line_indent = INDENT
    else:
        first.paragraph_format.first_line_indent = None
    for text in paragraphs[1:]:
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = FONT_SIZE
        if indent:
            p.paragraph_format.first_line_indent = INDENT
        else:
            p.paragraph_format.first_line_indent = None

def fill_schedule(cell, items):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for text in items:
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = FONT_SIZE

def main():
    tmpl = os.path.join("/Volumes/Samsung1T", "毕设相关文件", "附件2.毕业设计（论文）任务书.docx")
    out = os.path.join("/Volumes/Samsung1T", "毕设相关文件", "2211231001_蔡万鑫_华侨大学毕业设计（论文）任务书.docx")
    with open(CONTENT, encoding="utf-8") as f:
        data = json.load(f)
    doc = Document(tmpl)
    t = doc.tables[2]
    print("Section 1: Purpose...")
    fill_cell(t.cell(1, 0), data["section_1"], indent=True)
    print("Section 2: Content and Requirements...")
    fill_cell(t.cell(3, 0), data["section_2"], indent=False)
    print("Section 3: References...")
    fill_cell(t.cell(5, 0), data["section_3"], indent=False)
    print("Section 4: Schedule...")
    fill_schedule(t.cell(6, 0), data["section_4"])
    doc.save(out)
    print(f"Done! Output: {out}")

if __name__ == "__main__":
    main()
